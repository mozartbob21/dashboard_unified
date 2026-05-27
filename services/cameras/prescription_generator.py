from pathlib import Path
from datetime import datetime
from zipfile import ZipFile
from docx import Document
import json
import os
import sys
import traceback


# services/cameras/prescription_generator.py -> корень проекта
BASE_DIR = Path(__file__).resolve().parents[2]

DATA_DIR = BASE_DIR / "data"
CAMERAS_STATE_FILE = DATA_DIR / "cameras" / "state" / "dashboard_state.json"

TEMPLATES_DIR = BASE_DIR / "templates" / "prescriptions"
GENERATED_DIR = BASE_DIR / "generated" / "prescriptions"

NO_VIDEO_STREAM_TEMPLATE_PATH = TEMPLATES_DIR / "no_video_stream.docx"
NO_CAMERA_TEMPLATE_PATH = TEMPLATES_DIR / "no_camera.docx"

GENERATION_RESULT_FILE = GENERATED_DIR / "generation_result.json"

GENERATED_DIR.mkdir(parents=True, exist_ok=True)


STATUS_TO_TEMPLATE = {
    "not_working": NO_VIDEO_STREAM_TEMPLATE_PATH,
    "not_connected": NO_CAMERA_TEMPLATE_PATH,
}


STATUS_TO_LABEL = {
    "working": "Работает",
    "not_working": "Нет видеопотока",
    "not_connected": "Камера отсутствует",
    "unknown": "Не определено",
}


def load_json_file(path: Path, default=None):
    if not path.exists():
        return default

    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"Ошибка чтения JSON {path}: {e}")
        return default


def save_json_file(path: Path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )


def extract_camera_items(state) -> list[dict]:
    """
    Достаёт список камер из dashboard_state.json.

    Поддерживает форматы:
    [
      {...},
      {...}
    ]

    или:

    {
      "items": [...]
    }

    или:

    {
      "rows": [...]
    }

    или:

    {
      "results": [...]
    }

    или:

    {
      "cameras": [...]
    }
    """

    if not state:
        return []

    if isinstance(state, list):
        return [item for item in state if isinstance(item, dict)]

    if not isinstance(state, dict):
        return []

    for key in (
        "items",
        "rows",
        "results",
        "cameras",
        "data",
        "records",
    ):
        value = state.get(key)

        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]

    return []


def normalize_text(value) -> str:
    return str(value or "").strip()


def normalize_camera_status(item: dict) -> str:
    raw_status = normalize_text(
        item.get("camera_status")
        or item.get("status")
        or item.get("stream_status")
        or item.get("check_status")
        or item.get("state")
    ).lower()

    if raw_status in (
        "working",
        "ok",
        "online",
        "success",
        "active",
        "green",
        "норма",
        "работает",
        "активна",
        "активный",
    ):
        return "working"

    if raw_status in (
        "not_working",
        "critical",
        "offline",
        "error",
        "failed",
        "fail",
        "broken",
        "red",
        "критично",
        "не работает",
        "неработает",
        "ошибка",
        "отключена",
        "офлайн",
    ):
        return "not_working"

    if raw_status in (
        "not_connected",
        "no_stream",
        "missing_stream",
        "no_url",
        "empty_url",
        "не подключена",
        "не подключен",
        "нет ссылки",
        "нет потока",
        "без ссылки",
        "камера отсутствует",
        "отсутствует",
    ):
        return "not_connected"

    online = item.get("online")
    available = item.get("available")
    has_stream = item.get("has_stream")
    stream_url = normalize_text(item.get("stream_url") or item.get("link_url"))

    if online is False or available is False:
        return "not_working"

    if has_stream is False:
        return "not_connected"

    if online is True or available is True:
        return "working"

    if not stream_url:
        return "not_connected"

    return "unknown"


def normalize_camera_item(item: dict) -> dict:
    if not isinstance(item, dict):
        item = {}

    result = dict(item)

    result["camera_status"] = normalize_camera_status(result)

    if not result.get("owner"):
        result["owner"] = (
            result.get("responsible")
            or result.get("organization")
            or result.get("contractor")
            or "Не указана организация"
        )

    if not result.get("address"):
        result["address"] = result.get("object_address") or result.get("object") or "Адрес не указан"

    if not result.get("city"):
        result["city"] = result.get("municipality", "")

    if not result.get("checked_at"):
        result["checked_at"] = result.get("last_check") or result.get("created_at") or "—"

    return result


def sanitize_filename(value: str) -> str:
    value = str(value or "").strip()

    bad_chars = '\\/:*?"<>|'
    for ch in bad_chars:
        value = value.replace(ch, "_")

    value = value.replace("\n", " ").replace("\r", " ")
    value = "_".join(value.split())

    if not value:
        value = "empty"

    return value[:120]


def build_context(item: dict) -> dict:
    owner = item.get("owner", "") or "Не указана организация"
    address = item.get("address", "") or "Адрес не указан"
    checked_at = item.get("checked_at", "") or "—"
    camera_status = item.get("camera_status", "") or "unknown"

    now = datetime.now()

    return {
        "owner": owner,
        "organization": owner,
        "address": address,
        "checked_at": checked_at,
        "date": now.strftime("%d.%m.%Y"),
        "datetime": now.strftime("%d.%m.%Y %H:%M:%S"),
        "camera_status": camera_status,
        "camera_status_label": STATUS_TO_LABEL.get(camera_status, camera_status),
        "municipality": item.get("municipality", "") or "",
        "city": item.get("city", "") or item.get("municipality", "") or "",
        "contract": item.get("contract", "") or "",
        "work_type": item.get("work_type", "") or "",
        "responsible": item.get("responsible", "") or "",
        "last_check": item.get("last_check", "") or "",
        "stream_url": item.get("stream_url", "") or item.get("link_url", "") or "",
        "id": item.get("id", "") or "",
    }


def replace_text_in_paragraph(paragraph, context: dict):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in context.items():
        value = str(value)

        variants = [
            "{{" + key + "}}",
            "{{ " + key + " }}",
            "{{" + key.upper() + "}}",
            "{{ " + key.upper() + " }}",
            "{" + key + "}",
            "{" + key.upper() + "}",
        ]

        for placeholder in variants:
            new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = new_text


def replace_text_in_table(table, context: dict):
    for row in table.rows:
        for cell in row.cells:
            replace_text_in_container(cell, context)


def replace_text_in_container(container, context: dict):
    for paragraph in container.paragraphs:
        replace_text_in_paragraph(paragraph, context)

    for table in container.tables:
        replace_text_in_table(table, context)


def replace_text_in_headers_and_footers(doc: Document, context: dict):
    for section in doc.sections:
        replace_text_in_container(section.header, context)
        replace_text_in_container(section.footer, context)


def fill_docx_template(template_path: Path, output_path: Path, context: dict) -> str:
    if not template_path.exists():
        raise FileNotFoundError(f"Не найден шаблон предписания: {template_path}")

    doc = Document(template_path)

    replace_text_in_container(doc, context)
    replace_text_in_headers_and_footers(doc, context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return str(output_path)


def get_template_path_for_item(item: dict) -> Path | None:
    camera_status = item.get("camera_status", "") or "unknown"
    return STATUS_TO_TEMPLATE.get(camera_status)


def build_prescription_filename(item: dict, index: int | None = None) -> str:
    context = build_context(item)

    camera_status = context["camera_status"]
    owner = sanitize_filename(context["owner"])
    address = sanitize_filename(context["address"])

    prefix_map = {
        "not_working": "no_video_stream",
        "not_connected": "no_camera",
    }

    prefix = prefix_map.get(camera_status, "prescription")

    parts = []

    if index is not None:
        parts.append(str(index).zfill(3))

    parts.extend([prefix, owner, address])

    filename = "_".join(parts) + ".docx"

    return filename[:180]


def generate_prescription_from_template(item: dict, index: int | None = None) -> str:
    template_path = get_template_path_for_item(item)

    if template_path is None:
        camera_status = item.get("camera_status", "") or "unknown"
        raise ValueError(f"Для статуса камеры '{camera_status}' не настроен шаблон предписания")

    filename = build_prescription_filename(item, index=index)
    output_path = GENERATED_DIR / filename

    context = build_context(item)

    return fill_docx_template(
        template_path=template_path,
        output_path=output_path,
        context=context,
    )


def filter_items_for_prescriptions(items: list[dict]) -> list[dict]:
    result = []

    for item in items:
        normalized_item = normalize_camera_item(item)
        camera_status = normalized_item.get("camera_status", "") or "unknown"

        if camera_status in STATUS_TO_TEMPLATE:
            result.append(normalized_item)

    return result


def generate_individual_prescriptions(items: list[dict]) -> list[dict]:
    prescription_items = filter_items_for_prescriptions(items)
    generated_items = []

    for index, item in enumerate(prescription_items, start=1):
        owner = item.get("owner", "") or "Не указана организация"
        address = item.get("address", "") or "Адрес не указан"
        checked_at = item.get("checked_at", "") or "—"
        camera_status = item.get("camera_status", "") or "unknown"

        file_path = generate_prescription_from_template(item, index=index)

        generated_items.append(
            {
                "owner": owner,
                "address": address,
                "checked_at": checked_at,
                "camera_status": camera_status,
                "camera_status_label": STATUS_TO_LABEL.get(camera_status, camera_status),
                "filename": Path(file_path).name,
                "file_path": file_path,
                "url": f"/generated/prescriptions/{Path(file_path).name}",
            }
        )

    return generated_items


def generate_combined_prescription(items: list[dict]) -> str:
    prescription_items = filter_items_for_prescriptions(items)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"combined_prescriptions_{timestamp}.docx"
    file_path = GENERATED_DIR / filename

    doc = Document()
    doc.add_heading("СВОДНОЕ ПРЕДПИСАНИЕ", level=0)
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    doc.add_paragraph(f"Количество проблемных адресов: {len(prescription_items)}")
    doc.add_paragraph("")

    if not prescription_items:
        doc.add_paragraph("Проблемных адресов для формирования предписаний не найдено.")
    else:
        for index, item in enumerate(prescription_items, start=1):
            owner = item.get("owner", "") or "Не указана организация"
            address = item.get("address", "") or "Адрес не указан"
            checked_at = item.get("checked_at", "") or "—"
            camera_status = item.get("camera_status", "") or "unknown"
            status_label = STATUS_TO_LABEL.get(camera_status, camera_status)

            doc.add_heading(f"{index}. {owner}", level=1)
            doc.add_paragraph(f"Адрес объекта: {address}")
            doc.add_paragraph(f"Муниципалитет: {item.get('municipality', '') or '—'}")
            doc.add_paragraph(f"Время проверки: {checked_at}")
            doc.add_paragraph(f"Выявленное нарушение: {status_label}")

            contract = item.get("contract", "") or ""
            work_type = item.get("work_type", "") or ""
            responsible = item.get("responsible", "") or ""
            stream_url = item.get("stream_url", "") or item.get("link_url", "") or ""

            if contract:
                doc.add_paragraph(f"Договор: {contract}")

            if work_type:
                doc.add_paragraph(f"Вид работ: {work_type}")

            if responsible:
                doc.add_paragraph(f"Ответственный: {responsible}")

            if stream_url:
                doc.add_paragraph(f"Ссылка на видеопоток: {stream_url}")

            if camera_status == "not_working":
                doc.add_paragraph(
                    "В ходе автоматической проверки системы видеонаблюдения выявлено "
                    "отсутствие видеопотока либо неработоспособность трансляции по указанному адресу."
                )
            elif camera_status == "not_connected":
                doc.add_paragraph(
                    "В ходе автоматической проверки выявлено отсутствие подключенной камеры "
                    "видеонаблюдения по указанному адресу."
                )

            doc.add_paragraph(
                "Необходимо устранить выявленное нарушение и обеспечить работоспособность "
                "системы видеонаблюдения."
            )
            doc.add_paragraph("")

    doc.save(file_path)

    return str(file_path)


def generate_zip_with_prescriptions(items: list[dict]) -> str:
    prescription_items = filter_items_for_prescriptions(items)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_filename = f"prescriptions_{timestamp}.zip"
    zip_path = GENERATED_DIR / zip_filename

    generated_files = []

    for index, item in enumerate(prescription_items, start=1):
        file_path = generate_prescription_from_template(item, index=index)
        generated_files.append(Path(file_path))

    with ZipFile(zip_path, "w") as zip_file:
        for file_path in generated_files:
            zip_file.write(file_path, arcname=file_path.name)

    return str(zip_path)


def generate_demo_prescription(address: str, owner: str, checked_at: str | None = None) -> str:
    item = {
        "owner": owner,
        "address": address,
        "checked_at": checked_at or "—",
        "camera_status": "not_working",
    }

    return generate_prescription_from_template(item)


def build_generation_summary(
    mode: str,
    all_items: list[dict],
    problem_items: list[dict],
    individual_files: list[dict] | None = None,
    combined_file: str | None = None,
    zip_file: str | None = None,
    errors: list[str] | None = None,
) -> dict:
    individual_files = individual_files or []
    errors = errors or []

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    return {
        "ok": not errors,
        "mode": mode,
        "generated_at": generated_at,
        "total_count": len(all_items),
        "problem_count": len(problem_items),
        "generated_count": len(individual_files),
        "individual_files": individual_files,
        "combined_file": combined_file,
        "combined_filename": Path(combined_file).name if combined_file else None,
        "combined_url": f"/generated/prescriptions/{Path(combined_file).name}" if combined_file else None,
        "zip_file": zip_file,
        "zip_filename": Path(zip_file).name if zip_file else None,
        "zip_url": f"/generated/prescriptions/{Path(zip_file).name}" if zip_file else None,
        "errors": errors,
    }


def run_generation(mode: str = "all") -> dict:
    print("STAGE: Чтение результатов проверки камер")

    state = load_json_file(CAMERAS_STATE_FILE, default=[])

    raw_items = extract_camera_items(state)
    all_items = [normalize_camera_item(item) for item in raw_items]
    problem_items = filter_items_for_prescriptions(all_items)

    print(f"Найдено камер всего: {len(all_items)}")
    print(f"Проблемных камер для предписаний: {len(problem_items)}")

    if not problem_items:
        summary = build_generation_summary(
            mode=mode,
            all_items=all_items,
            problem_items=problem_items,
            errors=[],
        )
        save_json_file(GENERATION_RESULT_FILE, summary)
        print("Проблемных камер не найдено.")
        return summary

    individual_files = []
    combined_file = None
    zip_file = None
    errors = []

    normalized_mode = normalize_text(mode).lower()

    try:
        if normalized_mode in ("individual", "all"):
            print("STAGE: Формирование отдельных предписаний")
            individual_files = generate_individual_prescriptions(problem_items)
            print(f"Сформировано отдельных документов: {len(individual_files)}")

        if normalized_mode in ("combined", "all"):
            print("STAGE: Формирование сводного предписания")
            combined_file = generate_combined_prescription(problem_items)
            print(f"Сводный документ: {combined_file}")

        if normalized_mode in ("zip", "archive", "all"):
            print("STAGE: Формирование ZIP-архива")
            zip_file = generate_zip_with_prescriptions(problem_items)
            print(f"ZIP-архив: {zip_file}")

        if normalized_mode not in ("individual", "combined", "zip", "archive", "all"):
            errors.append(f"Неизвестный режим формирования: {mode}")

    except Exception as e:
        error_text = str(e)
        errors.append(error_text)
        print(f"Ошибка формирования предписаний: {error_text}")
        traceback.print_exc()

    summary = build_generation_summary(
        mode=mode,
        all_items=all_items,
        problem_items=problem_items,
        individual_files=individual_files,
        combined_file=combined_file,
        zip_file=zip_file,
        errors=errors,
    )

    save_json_file(GENERATION_RESULT_FILE, summary)

    print("STAGE: Сохранение результата")
    print(f"Результат сохранён: {GENERATION_RESULT_FILE}")

    return summary


def main():
    """
    Запуск из app.py:

    python -m services.cameras.prescription_generator

    По умолчанию формирует всё:
    - отдельные docx;
    - общий docx;
    - zip.

    Можно задать режим через переменную окружения:

    CAMERA_PRESCRIPTION_MODE=individual
    CAMERA_PRESCRIPTION_MODE=combined
    CAMERA_PRESCRIPTION_MODE=zip
    CAMERA_PRESCRIPTION_MODE=all
    """

    mode = os.getenv("CAMERA_PRESCRIPTION_MODE", "all")

    print("STAGE: Запуск формирования предписаний")
    print(f"BASE_DIR: {BASE_DIR}")
    print(f"CAMERAS_STATE_FILE: {CAMERAS_STATE_FILE}")
    print(f"TEMPLATES_DIR: {TEMPLATES_DIR}")
    print(f"GENERATED_DIR: {GENERATED_DIR}")
    print(f"MODE: {mode}")

    summary = run_generation(mode=mode)

    if summary.get("errors"):
        print("STAGE: Ошибка")
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        sys.exit(1)

    print("STAGE: Готово")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()